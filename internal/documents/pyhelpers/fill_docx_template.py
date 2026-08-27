#!/usr/bin/env python3
"""Fill {{placeholder}} tokens in an existing docx via python-docx,
preserving every paragraph's formatting except paragraphs that actually
contain a replaced placeholder (see replace_in_paragraph below for the
documented trade-off there). Prints one JSON status object to stdout:
{"replacements_applied": N} — N counts paragraphs touched, not
individual tokens, so a paragraph with two placeholders in it still
counts once.

Usage: fill_docx_template.py <template.docx> <output.docx> <replacements.json>
replacements.json is a flat {"name": "value", ...} object; "value" is
inserted as literal text (python-docx handles XML escaping internally
when a run's .text is set — no manual escaping needed here, unlike a
raw-XML-patching approach would require).

Exit 1 with a message on stderr for: python-docx not installed, the
template file missing/unreadable/not-a-valid-docx, or malformed
replacements JSON.
"""
import json
import sys


def replace_in_paragraph(paragraph, replacements):
    """Replace every {{name}} token in paragraph's text with its value.

    Trade-off, not a bug: this collapses the WHOLE paragraph to the
    formatting of its first run once any replacement fires in it —
    Word commonly splits one visually-contiguous word across multiple
    <w:r> runs (spell-check, autocorrect, manual formatting changes),
    so matching a token that might span run boundaries means working
    on the paragraph's full concatenated text, not run-by-run. A
    paragraph containing a placeholder that ALSO has other runs with
    different formatting (e.g. one bolded word elsewhere in the same
    sentence) loses that variation. Paragraphs with no placeholder are
    completely untouched.
    """
    if not paragraph.runs:
        return False
    full_text = "".join(run.text for run in paragraph.runs)
    if not any("{{" + name + "}}" in full_text for name in replacements):
        return False
    new_text = full_text
    for name, value in replacements.items():
        new_text = new_text.replace("{{" + name + "}}", value)
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def walk(paragraphs, replacements):
    return sum(1 for p in paragraphs if replace_in_paragraph(p, replacements))


def main():
    if len(sys.argv) != 4:
        print("usage: fill_docx_template.py <template.docx> <output.docx> <replacements.json>", file=sys.stderr)
        return 1
    template_path, output_path, replacements_path = sys.argv[1], sys.argv[2], sys.argv[3]

    try:
        from docx import Document
    except ImportError as exc:
        print(f"python-docx not installed: {exc}", file=sys.stderr)
        return 1

    try:
        with open(replacements_path, "r", encoding="utf-8") as f:
            replacements = json.load(f)
    except (OSError, json.JSONDecodeError) as exc:
        print(f"could not read replacements JSON {replacements_path!r}: {exc}", file=sys.stderr)
        return 1
    if not isinstance(replacements, dict):
        print(f"replacements JSON must be a flat object, got {type(replacements).__name__}", file=sys.stderr)
        return 1

    try:
        doc = Document(template_path)
    except Exception as exc:  # noqa: BLE001 - surfacing any open failure verbatim is the point
        print(f"could not open template {template_path!r}: {exc}", file=sys.stderr)
        return 1

    applied = walk(doc.paragraphs, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                applied += walk(cell.paragraphs, replacements)
    for section in doc.sections:
        applied += walk(section.header.paragraphs, replacements)
        applied += walk(section.footer.paragraphs, replacements)

    try:
        doc.save(output_path)
    except OSError as exc:
        print(f"could not save {output_path!r}: {exc}", file=sys.stderr)
        return 1

    print(json.dumps({"replacements_applied": applied}))
    return 0


if __name__ == "__main__":
    sys.exit(main())
